"""
Build an IEEE-style research paper for the study-hours prediction project.

The paper keeps the language simple, but it follows a formal technical paper
structure. All reported numbers come from the verified project data.
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_FILE = PROJECT_ROOT / "Study_Hours_Marks_Prediction_IEEE_Research_Paper.docx"
GRAPH_FILE = PROJECT_ROOT / "images" / "study_hours_marks_graph.png"

PAGE_WIDTH_DXA = 12240
LEFT_MARGIN_DXA = 900
RIGHT_MARGIN_DXA = 900
COLUMN_GAP_DXA = 360
COLUMN_WIDTH_DXA = (
    PAGE_WIDTH_DXA - LEFT_MARGIN_DXA - RIGHT_MARGIN_DXA - COLUMN_GAP_DXA
) // 2


def set_font(run, size=10, bold=False, italic=False):
    """Apply the Times New Roman font used by IEEE paper templates."""

    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_columns(section, number=2, space_dxa=COLUMN_GAP_DXA):
    """Set the number of text columns for a Word section."""

    section_properties = section._sectPr
    columns = section_properties.find(qn("w:cols"))

    if columns is None:
        columns = OxmlElement("w:cols")
        section_properties.append(columns)

    columns.set(qn("w:num"), str(number))
    columns.set(qn("w:space"), str(space_dxa))
    columns.set(qn("w:equalWidth"), "1")


def set_cell_margins(cell, top=60, start=70, bottom=60, end=70):
    """Add small but clear space inside a compact IEEE table."""

    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.find(qn("w:tcMar"))

    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)

    for side, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        margin = margins.find(qn(f"w:{side}"))
        if margin is None:
            margin = OxmlElement(f"w:{side}")
            margins.append(margin)
        margin.set(qn("w:type"), "dxa")
        margin.set(qn("w:w"), str(value))


def set_repeat_header(row):
    """Make a table header repeat after a page or column break."""

    row_properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    row_properties.append(repeat)


def apply_exact_table_geometry(table, weights, total_width=COLUMN_WIDTH_DXA):
    """Make table width, grid width, and cell widths match exactly."""

    weight_sum = sum(weights)
    widths = [round(total_width * weight / weight_sum) for weight in weights]
    widths[-1] += total_width - sum(widths)

    table_properties = table._tbl.tblPr

    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(total_width))

    layout = table_properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    indent = table_properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "0")

    table_grid = table._tbl.tblGrid
    for grid_column in list(table_grid):
        table_grid.remove(grid_column)

    for width in widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        table_grid.append(grid_column)

    for row in table.rows:
        for column_index, cell in enumerate(row.cells):
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(widths[column_index]))


def set_repeatable_paragraph_rules(paragraph):
    """Avoid isolated first or last lines where Word supports the rule."""

    paragraph_properties = paragraph._p.get_or_add_pPr()

    for tag in ("w:widowControl",):
        element = paragraph_properties.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            paragraph_properties.append(element)


def configure_document(document):
    """Set the page and default text rules."""

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    heading_one = document.styles["Heading 1"]
    heading_one.font.name = "Times New Roman"
    heading_one._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    heading_one._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    heading_one.font.size = Pt(10)
    heading_one.font.bold = False
    heading_one.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_one.paragraph_format.space_before = Pt(7)
    heading_one.paragraph_format.space_after = Pt(3)
    heading_one.paragraph_format.keep_with_next = True

    heading_two = document.styles["Heading 2"]
    heading_two.font.name = "Times New Roman"
    heading_two._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    heading_two._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    heading_two.font.size = Pt(10)
    heading_two.font.bold = False
    heading_two.font.italic = True
    heading_two.paragraph_format.space_before = Pt(5)
    heading_two.paragraph_format.space_after = Pt(1)
    heading_two.paragraph_format.keep_with_next = True

    first_section = document.sections[0]
    first_section.page_width = Inches(8.5)
    first_section.page_height = Inches(11)
    first_section.top_margin = Inches(0.75)
    first_section.bottom_margin = Inches(1)
    first_section.left_margin = Inches(0.625)
    first_section.right_margin = Inches(0.625)
    first_section.header_distance = Inches(0.3)
    first_section.footer_distance = Inches(0.3)
    set_columns(first_section, number=1)


def add_title_block(document):
    """Add the full-width IEEE title and author block."""

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run(
        "Study Hours and Marks Prediction Using Linear Regression: "
        "A Reproducible Beginner Study"
    )
    set_font(title_run, size=20)

    author = document.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(2)
    author_run = author.add_run("Suvodeep Roy")
    set_font(author_run, size=11)

    affiliation = document.add_paragraph()
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation.paragraph_format.space_after = Pt(10)
    affiliation_run = affiliation.add_run("Independent Student Researcher")
    set_font(affiliation_run, size=10, italic=True)

    two_column_section = document.add_section(WD_SECTION.CONTINUOUS)
    two_column_section.page_width = Inches(8.5)
    two_column_section.page_height = Inches(11)
    two_column_section.top_margin = Inches(0.75)
    two_column_section.bottom_margin = Inches(1)
    two_column_section.left_margin = Inches(0.625)
    two_column_section.right_margin = Inches(0.625)
    set_columns(two_column_section, number=2)


def add_body(document, text, first_line=True):
    """Add one justified IEEE body paragraph."""

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.first_line_indent = (
        Inches(0.14) if first_line else Inches(0)
    )
    set_repeatable_paragraph_rules(paragraph)
    run = paragraph.add_run(text)
    set_font(run, size=10)
    return paragraph


def add_label_paragraph(document, label, text):
    """Add IEEE abstract or index-term text."""

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(5)

    label_run = paragraph.add_run(label)
    set_font(label_run, size=9, bold=True, italic=True)

    text_run = paragraph.add_run(text)
    set_font(text_run, size=9, italic=True)
    return paragraph


def add_section_heading(document, text):
    """Add an IEEE main section heading."""

    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text.upper())
    set_font(run, size=10)
    run.font.small_caps = True
    return paragraph


def add_subheading(document, text):
    """Add an IEEE subsection heading."""

    paragraph = document.add_paragraph(style="Heading 2")
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_font(run, size=10, italic=True)
    return paragraph


def add_equation(document, equation_text, equation_number):
    """Add a centered equation with a visible IEEE equation number."""

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"{equation_text}       ({equation_number})")
    set_font(run, size=10, italic=True)


def add_caption(document, text, above=True):
    """Add a compact IEEE table or figure caption."""

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4 if above else 2)
    paragraph.paragraph_format.space_after = Pt(2 if above else 4)
    paragraph.paragraph_format.keep_with_next = above
    run = paragraph.add_run(text)
    set_font(run, size=8)
    run.font.small_caps = above
    return paragraph


def add_compact_table(document, headers, rows, weights):
    """Add a narrow table that fits one IEEE text column."""

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    header_row = table.rows[0]
    set_repeat_header(header_row)

    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(paragraph.runs[0], size=7.5, bold=True)

    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[index])
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if len(str(value)) <= 12
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            set_font(paragraph.runs[0], size=7.5)

    apply_exact_table_geometry(table, weights)
    return table


def add_abstract(document):
    """Add the IEEE abstract and index terms."""

    add_label_paragraph(
        document,
        "Abstract—",
        "This paper presents a small and reproducible machine learning study that "
        "predicts marks from study hours. The work uses a synthetic data set of "
        "20 records and a simple Linear Regression model. A fixed 75:25 train-test "
        "split gives an R-squared value of 0.997, a Mean Absolute Error of 1.69 "
        "marks, and a Root Mean Squared Error of 1.86 marks. A mean-value baseline "
        "performs much worse, with an R-squared value of -0.059 and a Mean Absolute "
        "Error of 30.64 marks. Five-fold cross-validation gives a mean R-squared "
        "value of 0.989 and a mean Mean Absolute Error of 1.83 marks. The program "
        "also checks invalid input, keeps predicted marks between 0 and 100, and "
        "includes ten unit tests. The results show that Linear Regression learns "
        "the clear line pattern in this sample data. They do not show that study "
        "hours alone can predict marks for real students. The main value of the "
        "work is a clear, safe, and beginner-friendly example of a full regression "
        "process."
    )
    add_label_paragraph(
        document,
        "Index Terms—",
        "educational data mining, Linear Regression, machine learning, marks "
        "prediction, reproducible study, study hours."
    )


def add_paper_content(document):
    """Add the complete IEEE paper content."""

    add_section_heading(document, "I. Introduction")
    add_body(
        document,
        "Educational data mining uses data and computer methods to study learning "
        "and student results. It can support tasks such as finding students who may "
        "need help, studying learning habits, and testing prediction methods [1]. "
        "Student performance prediction has been studied with decision trees, "
        "regression methods, and deep learning models [2]-[5]."
    )
    add_body(
        document,
        "Many published studies use large real data sets with many input fields. "
        "Such studies are useful, but they can be hard for a first-time learner to "
        "follow. This paper studies a much smaller question: can a single-input "
        "Linear Regression model learn a clear relation between study hours and "
        "marks in a synthetic teaching data set?"
    )
    add_body(
        document,
        "The contribution is educational rather than a new prediction method. The "
        "paper gives a full and repeatable process, reports a baseline, uses a fixed "
        "test split and five-fold cross-validation, states all limits, and links the "
        "reported results to tested source code."
    )

    add_section_heading(document, "II. Related Work")
    add_body(
        document,
        "Romero and Ventura described educational data mining as a broad field that "
        "includes prediction, grouping, relation study, and data display [1]. Cortez "
        "and Silva used school, social, grade, and study data to model student results "
        "with both regression and classification methods [2]. Their public data set "
        "contains hundreds of records and many input fields [6]."
    )
    add_body(
        document,
        "Yadav and Pal used decision-tree methods to group engineering students by "
        "likely result [3]. Kim, Vizitei, and Ganapathi used deep learning to predict "
        "online course outcomes from student event sequences [4]. Abd Elrahman et al. "
        "compared several classification and regression methods using eTextbook use "
        "data [5]. These studies show that real performance prediction often needs "
        "many records, many features, and careful model checks."
    )
    add_body(
        document,
        "The present study does not try to match those systems. It uses one feature "
        "and synthetic data so that a new learner can inspect each step. This clear "
        "scope is important because a simple teaching model must not be presented as "
        "a tool for real student decisions."
    )

    add_section_heading(document, "III. Research Design")
    add_subheading(document, "A. Research Questions")
    add_body(
        document,
        "The study asks three questions. First, can Linear Regression learn the line "
        "pattern in the synthetic data? Second, does it perform better than a simple "
        "mean-value baseline? Third, does its result stay stable across five data "
        "folds?"
    )

    add_subheading(document, "B. Data")
    add_body(
        document,
        "The data set contains 20 synthetic records. Study hours range from 1.0 to "
        "10.5, and marks range from 12 to 98. Small changes were added to the marks "
        "so that the data does not form a perfect line. No real person, school record, "
        "private value, or survey answer is used."
    )
    add_caption(document, "TABLE I\nSAMPLE DATA RECORDS")
    add_compact_table(
        document,
        ["Hours", "Marks", "Use"],
        [
            ("1.0", "12", "Sample"),
            ("3.0", "31", "Sample"),
            ("5.0", "54", "Sample"),
            ("7.0", "72", "Sample"),
            ("9.0", "88", "Sample"),
            ("10.5", "98", "Sample"),
        ],
        [1, 1, 1.5],
    )

    add_subheading(document, "C. Model")
    add_body(
        document,
        "Linear Regression estimates a straight line between an input x and an "
        "output y. In this study, x is study hours and y is marks. The model is"
    )
    add_equation(document, "y_hat = b_0 + b_1 x", "1")
    add_body(
        document,
        "where y_hat is the predicted mark, b_0 is the line start, and b_1 is the "
        "line slope. The implementation uses the Scikit-learn LinearRegression "
        "class [7]. Pandas stores the data [8], and Matplotlib creates the graph [9]."
    )

    add_subheading(document, "D. Test Process")
    add_body(
        document,
        "The main test keeps 25 percent of the records for testing and uses random "
        "state 42. The same test records are also used with a DummyRegressor that "
        "always predicts the mean training mark. This gives a clear baseline. A "
        "five-fold cross-validation test then trains and tests the Linear Regression "
        "model five times with different folds."
    )
    add_body(
        document,
        "The reported measures are R-squared, Mean Absolute Error, and Root Mean "
        "Squared Error. They are defined as"
    )
    add_equation(document, "R^2 = 1 - SS_res / SS_tot", "2")
    add_equation(document, "MAE = (1/n) sum |y_i - y_hat_i|", "3")
    add_equation(document, "RMSE = sqrt[(1/n) sum (y_i - y_hat_i)^2]", "4")
    add_body(
        document,
        "A higher R-squared value is better. Lower MAE and RMSE values are better. "
        "RMSE gives more weight to large errors."
    )

    add_section_heading(document, "IV. Implementation")
    add_body(
        document,
        "The program is written in Python. Each main task is placed in a small "
        "function. Separate functions create data, split data, train the model, "
        "evaluate the model, check user input, predict marks, and save the graph. "
        "The prediction input uses the same named Hours column used during training. "
        "This avoids the Scikit-learn feature-name warning."
    )
    add_body(
        document,
        "The input check rejects an empty value, text, a negative value, a value "
        "above 24, NaN, and infinity. The final mark is limited to the valid range "
        "from 0 to 100. Ten unit tests check the data shape, valid input, invalid "
        "input, model measures, and prediction range."
    )

    add_section_heading(document, "V. Results")
    add_caption(document, "TABLE II\nFIXED-SPLIT MODEL COMPARISON")
    add_compact_table(
        document,
        ["Model", "R-squared", "MAE", "RMSE"],
        [
            ("Linear Reg.", "0.997", "1.69", "1.86"),
            ("Mean base", "-0.059", "30.64", "33.97"),
        ],
        [1.5, 1, 0.8, 0.8],
    )
    add_body(
        document,
        "Table II shows that Linear Regression performs far better than the mean "
        "baseline on the fixed test split. The model equation learned from the "
        "training part is"
    )
    add_equation(document, "y_hat = 5.28 + 9.20x", "5")
    add_body(
        document,
        "For five study hours, this model predicts 51.30 marks. The five test marks "
        "have absolute errors from 0.72 to 2.70 marks."
    )

    add_caption(document, "TABLE III\nFIVE-FOLD CROSS-VALIDATION")
    add_compact_table(
        document,
        ["Measure", "Mean", "Std. Dev."],
        [
            ("R-squared", "0.989", "0.006"),
            ("MAE", "1.83", "0.19"),
            ("RMSE", "2.12", "0.37"),
        ],
        [1.6, 1, 1],
    )
    add_body(
        document,
        "The five-fold results remain strong across all folds. The mean R-squared "
        "value is 0.989. The small standard deviation shows low change across the "
        "five folds. This is expected because the synthetic data follows a clear "
        "near-linear pattern."
    )

    if GRAPH_FILE.exists():
        figure = document.add_paragraph()
        figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
        figure.paragraph_format.space_before = Pt(5)
        picture = figure.add_run().add_picture(
            str(GRAPH_FILE),
            width=Inches(3.25),
        )
        picture._inline.docPr.set(
            "descr",
            "Blue points show synthetic marks for study hours. "
            "A red line shows the Linear Regression model.",
        )
        picture._inline.docPr.set(
            "title",
            "Study hours and marks regression graph",
        )
        add_caption(
            document,
            "Fig. 1. Synthetic data points and the fitted regression line.",
            above=False,
        )

    add_section_heading(document, "VI. Discussion")
    add_body(
        document,
        "The fixed split, baseline comparison, and cross-validation all show the "
        "same result: a straight-line model matches this synthetic data very well. "
        "The baseline result is important because it shows that the model learns more "
        "than a simple average. Cross-validation reduces the risk that one favorable "
        "split alone caused the reported score."
    )
    add_body(
        document,
        "However, the high score is mainly a property of the data design. The data "
        "has only 20 records, one input, and a clear line pattern. Real marks can also "
        "depend on prior knowledge, class attendance, sleep, health, teaching quality, "
        "assessment type, and many social factors. The model therefore shows a method, "
        "not a general rule about students."
    )
    add_body(
        document,
        "The study also does not prove cause. It shows an association inside the "
        "sample data. It does not prove that adding one study hour will cause a fixed "
        "increase of 9.20 marks for a real student."
    )

    add_section_heading(document, "VII. Ethics, Safety, and Reproducibility")
    add_body(
        document,
        "No human data was collected, so this study creates no direct privacy risk. "
        "If the work is extended with real student records, informed consent, data "
        "minimization, secure storage, access control, and institutional approval may "
        "be needed. A prediction must not be used alone to rank, punish, or deny help "
        "to a student."
    )
    add_body(
        document,
        "The data, fixed random state, source code, tests, model settings, and exact "
        "reported measures are included in the project. This supports repeat work. "
        "The program was checked with Python syntax tests, dependency checks, ten unit "
        "tests, invalid-input tests, and graph generation."
    )

    add_section_heading(document, "VIII. Limitations and Future Work")
    add_body(
        document,
        "The main limits are the synthetic data, small sample size, single input, "
        "limited data range, and use of one model type. A future study should use a "
        "larger real data set with consent, define the target and study period before "
        "data collection, check missing data and bias, compare several models, use a "
        "held-out test set, and report confidence ranges."
    )
    add_body(
        document,
        "A suitable next step is the UCI Student Performance data set [6]. It has "
        "hundreds of records and includes study time, grades, school facts, and social "
        "facts. Such work would require stronger ethics, feature study, and model "
        "validation than the present teaching project."
    )

    add_section_heading(document, "IX. Conclusion")
    add_body(
        document,
        "This paper presented a simple and reproducible Linear Regression study for "
        "marks prediction. Linear Regression clearly outperformed a mean baseline and "
        "gave stable five-fold results on the synthetic data. The program also adds "
        "safe input checks, a bounded output, a graph, and unit tests."
    )
    add_body(
        document,
        "The work is professional as a teaching and portfolio study, but it is not a "
        "validated real-student prediction system. Its main result is a clear example "
        "of correct regression practice: define the scope, keep test data separate, "
        "compare with a baseline, report more than one measure, state limits, and avoid "
        "claims that the data cannot support."
    )

    add_section_heading(document, "References")
    references = [
        (
            "[1] C. Romero and S. Ventura, \"Educational data mining: A review "
            "of the state-of-the-art,\" IEEE Trans. Syst., Man, Cybern. C, "
            "vol. 40, no. 6, pp. 601-618, Nov. 2010."
        ),
        (
            "[2] P. Cortez and A. M. G. Silva, \"Using data mining to predict "
            "secondary school student performance,\" in Proc. 5th Future Bus. "
            "Technol. Conf., Porto, Portugal, 2008, pp. 5-12."
        ),
        (
            "[3] S. K. Yadav and S. Pal, \"Data mining: A prediction for "
            "performance improvement of engineering students using "
            "classification,\" arXiv:1203.3832, 2012, "
            "doi: 10.48550/arXiv.1203.3832."
        ),
        (
            "[4] B.-H. Kim, E. Vizitei, and V. Ganapathi, \"GritNet: Student "
            "performance prediction with deep learning,\" arXiv:1804.07405, "
            "2018, doi: 10.48550/arXiv.1804.07405."
        ),
        (
            "[5] A. Abd Elrahman, T. H. A. Soliman, A. I. Taloba, and M. F. "
            "Farghally, \"A predictive model for student performance in "
            "classrooms using student interactions with an eTextbook,\" "
            "arXiv:2203.03713, 2022, doi: 10.48550/arXiv.2203.03713."
        ),
        (
            "[6] P. Cortez, \"Student Performance,\" UCI Machine Learning "
            "Repository, 2008, doi: 10.24432/C5TG7T."
        ),
        (
            "[7] F. Pedregosa et al., \"Scikit-learn: Machine learning in "
            "Python,\" J. Mach. Learn. Res., vol. 12, pp. 2825-2830, 2011."
        ),
        (
            "[8] W. McKinney, \"Data structures for statistical computing in "
            "Python,\" in Proc. 9th Python Sci. Conf., 2010, pp. 56-61."
        ),
        (
            "[9] J. D. Hunter, \"Matplotlib: A 2D graphics environment,\" "
            "Comput. Sci. Eng., vol. 9, no. 3, pp. 90-95, 2007."
        ),
    ]

    for reference in references:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.first_line_indent = Inches(-0.18)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(reference)
        set_font(run, size=8)


def build_ieee_paper():
    """Create and save the final IEEE-style Word paper."""

    document = Document()
    configure_document(document)
    add_title_block(document)
    add_abstract(document)
    add_paper_content(document)

    document.core_properties.title = (
        "Study Hours and Marks Prediction Using Linear Regression: "
        "A Reproducible Beginner Study"
    )
    document.core_properties.author = "Suvodeep Roy"
    document.core_properties.subject = (
        "IEEE-style educational machine learning research paper"
    )
    document.core_properties.keywords = (
        "educational data mining, linear regression, marks prediction, "
        "study hours, reproducible machine learning"
    )

    document.save(OUTPUT_FILE)
    print(OUTPUT_FILE)


if __name__ == "__main__":
    build_ieee_paper()
