"""
Create a simple and beginner-friendly Word research paper.

This helper file creates the final DOCX paper from the verified project facts.
It is not part of the machine learning program.
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_FILE = PROJECT_ROOT / "Study_Hours_Marks_Prediction_Research_Paper.docx"
GRAPH_FILE = PROJECT_ROOT / "images" / "study_hours_marks_graph.png"

BLUE = "1F4E79"
LIGHT_BLUE = "DCE6F1"
LIGHT_GRAY = "F2F2F2"
DARK_GRAY = RGBColor(80, 80, 80)


def set_cell_shading(cell, fill_color):
    """Give a table cell a light background color."""

    cell_properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_color)
    cell_properties.append(shading)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    """Add space between table text and table borders."""

    cell_properties = cell._tc.get_or_add_tcPr()
    cell_margins = cell_properties.first_child_found_in("w:tcMar")

    if cell_margins is None:
        cell_margins = OxmlElement("w:tcMar")
        cell_properties.append(cell_margins)

    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        margin = cell_margins.find(qn(f"w:{margin_name}"))
        if margin is None:
            margin = OxmlElement(f"w:{margin_name}")
            cell_margins.append(margin)
        margin.set(qn("w:w"), str(margin_value))
        margin.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    """Repeat a table header when a table moves to another page."""

    row_properties = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    row_properties.append(table_header)


def apply_exact_table_widths(table, width_weights):
    """Set exact Word table, grid, and cell widths in DXA units."""

    total_width_dxa = 9360
    total_weight = sum(width_weights)
    column_widths = [
        round(total_width_dxa * width / total_weight)
        for width in width_weights
    ]
    column_widths[-1] += total_width_dxa - sum(column_widths)

    table_properties = table._tbl.tblPr

    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(total_width_dxa))

    table_layout = table_properties.find(qn("w:tblLayout"))
    if table_layout is None:
        table_layout = OxmlElement("w:tblLayout")
        table_properties.append(table_layout)
    table_layout.set(qn("w:type"), "fixed")

    table_indent = table_properties.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:type"), "dxa")
    table_indent.set(qn("w:w"), "0")

    table_grid = table._tbl.tblGrid
    for old_column in list(table_grid):
        table_grid.remove(old_column)

    for column_width in column_widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(column_width))
        table_grid.append(grid_column)

    for row in table.rows:
        for column_index, cell in enumerate(row.cells):
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(column_widths[column_index]))


def add_page_number(paragraph):
    """Add an automatic page number field."""

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    number_text = OxmlElement("w:t")
    number_text.text = "1"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(separate)
    run._r.append(number_text)
    run._r.append(end)


def set_run_font(run, name="Arial", size=12, bold=False, color=None):
    """Apply the same readable font in Word and other document tools."""

    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold

    if color:
        run.font.color.rgb = color


def configure_styles(document):
    """Create a clear style system for the whole paper."""

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    title_style = document.styles["Title"]
    title_style.font.name = "Arial"
    title_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    title_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(31, 78, 121)
    title_style.paragraph_format.space_after = Pt(16)

    subtitle_style = document.styles["Subtitle"]
    subtitle_style.font.name = "Arial"
    subtitle_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    subtitle_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    subtitle_style.font.size = Pt(13)
    subtitle_style.font.color.rgb = DARK_GRAY

    for style_name, size in [
        ("Heading 1", 16),
        ("Heading 2", 14),
        ("Heading 3", 12),
    ]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def configure_page(document):
    """Set page size, margins, header, and footer."""

    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "Study Hours vs Marks Prediction"
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header_paragraph.runs[0], size=9, color=DARK_GRAY)

    paragraph_properties = header_paragraph._p.get_or_add_pPr()
    border_group = OxmlElement("w:pBdr")
    bottom_border = OxmlElement("w:bottom")
    bottom_border.set(qn("w:val"), "single")
    bottom_border.set(qn("w:sz"), "4")
    bottom_border.set(qn("w:space"), "3")
    bottom_border.set(qn("w:color"), "B7B7B7")
    border_group.append(bottom_border)
    paragraph_properties.append(border_group)

    add_page_number(section.footer.paragraphs[0])
    for run in section.footer.paragraphs[0].runs:
        set_run_font(run, size=9, color=DARK_GRAY)


def add_body_paragraph(document, text, bold_start=None):
    """Add a justified body paragraph with an optional bold start."""

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if bold_start and text.startswith(bold_start):
        bold_run = paragraph.add_run(bold_start)
        set_run_font(bold_run, bold=True)
        normal_run = paragraph.add_run(text[len(bold_start):])
        set_run_font(normal_run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)

    return paragraph


def add_bullet(document, text):
    """Add a real Word bullet with simple spacing."""

    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    set_run_font(run)


def add_numbered_item(document, text):
    """Add a real Word numbered item."""

    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    set_run_font(run)


def add_table(document, headers, rows, widths=None):
    """Add a clean table with a light header and readable cell space."""

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    header_row = table.rows[0]
    set_repeat_table_header(header_row)

    for column_index, header_text in enumerate(headers):
        cell = header_row.cells[column_index]
        cell.text = header_text
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(paragraph.runs[0], size=10, bold=True)

    for row_values in rows:
        row_cells = table.add_row().cells

        for column_index, value in enumerate(row_values):
            cell = row_cells[column_index]
            cell.text = str(value)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if len(str(value)) < 18
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            set_run_font(paragraph.runs[0], size=10)

    apply_exact_table_widths(
        table,
        widths if widths else [1] * len(headers),
    )

    document.add_paragraph()
    return table


def add_cover_page(document):
    """Create a simple academic cover page."""

    document.add_paragraph()
    document.add_paragraph()

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(
        "Study Hours and Marks Prediction Using Linear Regression"
    )
    set_run_font(title_run, size=22, bold=True, color=RGBColor(31, 78, 121))

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "A Simple Machine Learning Study for New Learners"
    )
    set_run_font(subtitle_run, size=13, color=DARK_GRAY)

    document.add_paragraph()

    line = document.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line.add_run("Research Paper")
    set_run_font(line_run, size=15, bold=True)

    document.add_paragraph()
    document.add_paragraph()

    details = [
        ("Author", "Suvodeep Roy"),
        ("Project", "Study Hours vs Marks Prediction"),
        ("Field", "Machine Learning"),
        ("Method", "Linear Regression"),
        ("Date", "June 25, 2026"),
    ]
    add_table(document, ["Paper Detail", "Value"], details, widths=[2.0, 4.3])

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run(
        "This paper uses simple English so that a new learner can read it with ease."
    )
    set_run_font(note_run, size=11, color=DARK_GRAY)

    document.add_page_break()


def add_abstract(document):
    """Add the abstract and key words."""

    document.add_heading("Abstract", level=1)
    add_body_paragraph(
        document,
        "This paper explains a small machine learning project that predicts marks "
        "from study hours. The project uses Linear Regression. It uses 20 sample "
        "records. Each record has study hours and marks. The data is split into "
        "a training part and a test part. The model learns from the training part. "
        "The test part is then used to check the model. The model gives an "
        "R-squared score of 0.997. It gives a Mean Absolute Error of 1.69 marks "
        "and a Mean Squared Error of 3.47. For five study hours, the model gives "
        "a result of 51.30 marks. The result shows that the model can learn the "
        "clear line pattern in this sample data. The paper also explains an "
        "important limit. The data is small and made for study. It is not real "
        "school data. So, this work must not be used to judge a real student. "
        "The main aim is to help a new learner understand data, model training, "
        "model tests, input checks, and graphs."
    )
    add_body_paragraph(
        document,
        "Key words: study hours, marks, machine learning, Linear Regression, "
        "prediction, Python, beginner project.",
        bold_start="Key words:",
    )


def add_contents(document):
    """Add a clear paper map."""

    document.add_heading("Paper Contents", level=1)
    contents = [
        "1. Introduction",
        "2. Problem Statement",
        "3. Aim and Objectives",
        "4. Research Questions",
        "5. Basic Review of Ideas",
        "6. Research Method",
        "7. Project Design",
        "8. Data Used in the Study",
        "9. Model Training and Testing",
        "10. Results",
        "11. Discussion",
        "12. Program Safety and Input Checks",
        "13. Test Plan",
        "14. Limits of the Study",
        "15. Future Work",
        "16. Conclusion",
        "17. References",
        "Appendix A. Main Program Flow",
        "Appendix B. Simple Viva Questions",
    ]

    for item in contents:
        add_bullet(document, item)

    document.add_page_break()


def add_main_sections(document):
    """Add the main research paper text."""

    document.add_heading("1. Introduction", level=1)
    add_body_paragraph(
        document,
        "Machine learning helps a computer learn a pattern from data. It can then "
        "use the pattern to give a new result. A mark is a number. This means that "
        "mark prediction is a regression task. Regression is used when the result "
        "is a number, such as marks, price, age, or time."
    )
    add_body_paragraph(
        document,
        "This project studies one simple link: study hours and marks. Study hours "
        "are used as the input. Marks are used as the output. A Linear Regression "
        "model is used because it is easy to learn, easy to show on a graph, and "
        "easy to explain in a viva."
    )
    add_body_paragraph(
        document,
        "The project is made for a new learner. The code uses clear names, small "
        "functions, short notes, input checks, model tests, and a graph. It also "
        "avoids a false claim. More study may help marks, but marks can also change "
        "due to sleep, health, class skill, past study, stress, and many other facts."
    )

    document.add_heading("2. Problem Statement", level=1)
    add_body_paragraph(
        document,
        "A new machine learning learner may know Python but may not know how a full "
        "prediction task works. The learner needs a small project that shows each "
        "step in a clear order. The project must show how to make data, split data, "
        "train a model, test a model, take safe user input, make a prediction, and "
        "show a graph."
    )
    add_body_paragraph(
        document,
        "The main problem is to build a simple model that can estimate marks from "
        "study hours while keeping the code and result easy to understand."
    )

    document.add_heading("3. Aim and Objectives", level=1)
    add_body_paragraph(
        document,
        "The main aim is to build and study a small Linear Regression model for "
        "marks prediction."
    )
    objectives = [
        "Create a small sample data set with study hours and marks.",
        "Use study hours as the input and marks as the output.",
        "Split the data into training data and test data.",
        "Train a Linear Regression model.",
        "Check the model with R-squared, Mean Absolute Error, and Mean Squared Error.",
        "Take study hours from the user in a safe way.",
        "Keep the result from 0 to 100 marks.",
        "Show the data and model line on a graph.",
        "Write code that a new learner can read and explain.",
    ]
    for objective in objectives:
        add_bullet(document, objective)

    document.add_heading("4. Research Questions", level=1)
    questions = [
        "Can a simple Linear Regression model learn the pattern in the sample data?",
        "How close are the model results to the test marks?",
        "Can the program stop wrong study-hour input?",
        "Can the full work be explained in simple English?",
    ]
    for question in questions:
        add_numbered_item(document, question)

    document.add_heading("5. Basic Review of Ideas", level=1)
    document.add_heading("5.1 Machine Learning", level=2)
    add_body_paragraph(
        document,
        "Machine learning is a way to let a computer learn from examples. The "
        "computer is not given one fixed answer for each new case. It finds a "
        "pattern in old data and uses that pattern for a new case."
    )

    document.add_heading("5.2 Supervised Learning", level=2)
    add_body_paragraph(
        document,
        "This project uses supervised learning. In supervised learning, each old "
        "record has an input and a known output. Here, the input is study hours and "
        "the known output is marks."
    )

    document.add_heading("5.3 Regression", level=2)
    add_body_paragraph(
        document,
        "Regression is used when the output is a number. This project predicts a "
        "mark value, so regression is the right task type."
    )

    document.add_heading("5.4 Linear Regression", level=2)
    add_body_paragraph(
        document,
        "Linear Regression finds a straight line that stays as close as it can to "
        "the data points. The line is written in this simple form:"
    )
    equation = document.add_paragraph()
    equation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation_run = equation.add_run(
        "Predicted Marks = (Line Slope x Study Hours) + Line Start"
    )
    set_run_font(equation_run, size=12, bold=True, color=RGBColor(31, 78, 121))
    add_body_paragraph(
        document,
        "The line slope shows the change in predicted marks when study hours rise "
        "by one hour. The line start is the value where the line begins."
    )

    document.add_heading("5.5 Tools Used", level=2)
    tool_rows = [
        ("Python", "Runs the full program."),
        ("Pandas", "Stores the data in rows and columns."),
        ("Scikit-learn", "Splits data, trains the model, and checks the model."),
        ("Matplotlib", "Makes the data and model graph."),
        ("unittest", "Checks the main code rules."),
    ]
    add_table(document, ["Tool", "Use in This Work"], tool_rows, widths=[1.7, 4.6])

    document.add_heading("6. Research Method", level=1)
    add_body_paragraph(
        document,
        "This work uses a small test-based method. The steps are simple and can be "
        "done again on the same computer."
    )
    method_steps = [
        "Create 20 sample records.",
        "Place the records in a Pandas DataFrame.",
        "Select Hours as the input column.",
        "Select Marks as the output column.",
        "Keep 25 percent of the records for testing.",
        "Use random state 42 so the split stays the same.",
        "Train the Linear Regression model with the training records.",
        "Predict marks for the test records.",
        "Find R-squared, Mean Absolute Error, and Mean Squared Error.",
        "Ask the user for study hours.",
        "Check the input and make the final mark prediction.",
        "Save a graph of the data and the model line.",
    ]
    for step in method_steps:
        add_numbered_item(document, step)

    document.add_heading("7. Project Design", level=1)
    add_body_paragraph(
        document,
        "The program keeps each main task in a small function. This makes the code "
        "easy to read and test. It also helps a learner explain one task at a time."
    )
    design_rows = [
        ("create_student_dataset", "Creates the 20 sample records."),
        ("prepare_training_and_testing_data", "Makes the train and test parts."),
        ("train_prediction_model", "Builds and trains the model."),
        ("evaluate_prediction_model", "Finds the three model test values."),
        ("validate_study_hours", "Checks user input."),
        ("predict_student_marks", "Makes one safe mark prediction."),
        ("save_and_display_graph", "Saves and shows the graph."),
        ("main", "Runs all steps in the right order."),
    ]
    add_table(document, ["Function", "Simple Task"], design_rows, widths=[2.8, 3.5])

    document.add_heading("8. Data Used in the Study", level=1)
    add_body_paragraph(
        document,
        "The data has 20 sample records. It is made only for learning. It is not "
        "taken from a school, college, survey, or real student group. Small changes "
        "were added so the points do not form a fully exact line."
    )
    dataset_rows = [
        ("1", "1.0", "12"),
        ("2", "1.5", "18"),
        ("3", "2.0", "21"),
        ("4", "2.5", "27"),
        ("5", "3.0", "31"),
        ("6", "3.5", "38"),
        ("7", "4.0", "43"),
        ("8", "4.5", "47"),
        ("9", "5.0", "54"),
        ("10", "5.5", "56"),
        ("11", "6.0", "63"),
        ("12", "6.5", "67"),
        ("13", "7.0", "72"),
        ("14", "7.5", "76"),
        ("15", "8.0", "81"),
        ("16", "8.5", "85"),
        ("17", "9.0", "88"),
        ("18", "9.5", "92"),
        ("19", "10.0", "95"),
        ("20", "10.5", "98"),
    ]
    add_table(
        document,
        ["Record", "Study Hours", "Marks"],
        dataset_rows,
        widths=[1.4, 2.4, 2.4],
    )
    add_body_paragraph(
        document,
        "The input range in the sample is 1.0 to 10.5 hours. The mark range is 12 "
        "to 98. A result far outside this range is less safe because the model has "
        "not learned from many such cases."
    )

    document.add_heading("9. Model Training and Testing", level=1)
    add_body_paragraph(
        document,
        "The program uses 75 percent of the records for training and 25 percent for "
        "testing. The model sees only the training part while it learns. The test "
        "part is used after training. This gives a fair basic check."
    )
    add_body_paragraph(
        document,
        "The fixed random state is 42. This does not make the model better. It only "
        "keeps the same split each time, which helps repeat the study."
    )

    document.add_heading("9.1 Test Values", level=2)
    result_rows = [
        ("12", "14.48", "2.48"),
        ("92", "92.72", "0.72"),
        ("85", "83.52", "1.48"),
        ("18", "19.09", "1.09"),
        ("54", "51.30", "2.70"),
    ]
    add_table(
        document,
        ["Real Marks", "Model Marks", "Gap"],
        result_rows,
        widths=[2.1, 2.1, 2.1],
    )

    document.add_heading("10. Results", level=1)
    metric_rows = [
        ("R-squared", "0.997", "The line fits this sample pattern very well."),
        ("Mean Absolute Error", "1.69 marks", "The mean gap is about 1.69 marks."),
        ("Mean Squared Error", "3.47", "Large gaps are low in this test set."),
        ("Five-hour result", "51.30 marks", "The model gives 51.30 marks for 5 hours."),
    ]
    add_table(
        document,
        ["Measure", "Result", "Simple Meaning"],
        metric_rows,
        widths=[1.8, 1.4, 3.1],
    )
    add_body_paragraph(
        document,
        "The model learned this line:"
    )
    learned_line = document.add_paragraph()
    learned_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    learned_run = learned_line.add_run(
        "Predicted Marks = (9.20 x Study Hours) + 5.28"
    )
    set_run_font(learned_run, size=12, bold=True, color=RGBColor(31, 78, 121))
    add_body_paragraph(
        document,
        "In this sample, one more study hour adds about 9.20 predicted marks. This "
        "is only the pattern of the sample. It is not a rule for all students."
    )

    document.add_heading("10.1 Graph Result", level=2)
    if GRAPH_FILE.exists():
        figure_paragraph = document.add_paragraph()
        figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        figure_run = figure_paragraph.add_run()
        picture = figure_run.add_picture(str(GRAPH_FILE), width=Inches(6.25))

        # Add a clear text description for people who use screen readers.
        picture_properties = picture._inline.docPr
        picture_properties.set(
            "descr",
            "Blue points show sample marks for study hours. "
            "A red line shows the Linear Regression result.",
        )
        picture_properties.set("title", "Study Hours and Marks Graph")

        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption.add_run(
            "Figure 1. Sample study-hour points and the Linear Regression line."
        )
        set_run_font(caption_run, size=10, color=DARK_GRAY)

    document.add_heading("11. Discussion", level=1)
    add_body_paragraph(
        document,
        "The R-squared score is close to 1 because the sample points have a very "
        "clear line pattern. The error values are also low. This means the model "
        "works well on this small sample."
    )
    add_body_paragraph(
        document,
        "The high score must be read with care. The data was made for this learning "
        "task. It has only one input and only 20 records. A real student data set "
        "will have more change and more facts. A high score here does not prove that "
        "the model will work with the same score in a real school."
    )
    add_body_paragraph(
        document,
        "The main value of the project is clear learning. It shows how data moves "
        "from a table to a model, from a model to a test, and from a test to a safe "
        "user result."
    )

    document.add_heading("12. Program Safety and Input Checks", level=1)
    add_body_paragraph(
        document,
        "The program checks each user value before prediction. This stops common "
        "errors and gives a clear message."
    )
    safety_rows = [
        ("Empty value", "Not allowed", "Study hours cannot be empty."),
        ("Text such as five", "Not allowed", "A number is required."),
        ("Less than zero", "Not allowed", "Hours cannot be negative."),
        ("More than 24", "Not allowed", "One day has only 24 hours."),
        ("NaN or infinity", "Not allowed", "A normal number is required."),
        ("Valid number", "Allowed", "The model makes a prediction."),
    ]
    add_table(
        document,
        ["Input Type", "Result", "Reason"],
        safety_rows,
        widths=[1.8, 1.3, 3.2],
    )
    add_body_paragraph(
        document,
        "The final mark is also kept from 0 to 100. This rule stops a line model "
        "from showing a mark below 0 or above 100."
    )

    document.add_heading("13. Test Plan", level=1)
    add_body_paragraph(
        document,
        "The project has ten unit tests. A unit test checks one small rule in the "
        "code. All ten tests passed on June 25, 2026."
    )
    test_rows = [
        ("Data columns", "Hours and Marks are present.", "Pass"),
        ("Data size", "There are 20 records.", "Pass"),
        ("Whole number", "The value 5 is allowed.", "Pass"),
        ("Decimal number", "The value 5.5 is allowed.", "Pass"),
        ("Empty value", "The value is stopped.", "Pass"),
        ("Text value", "The value is stopped.", "Pass"),
        ("Negative value", "The value is stopped.", "Pass"),
        ("Above 24", "The value is stopped.", "Pass"),
        ("NaN value", "The value is stopped.", "Pass"),
        ("Mark range", "The result stays from 0 to 100.", "Pass"),
    ]
    add_table(
        document,
        ["Test", "Expected Work", "Status"],
        test_rows,
        widths=[1.7, 3.5, 1.1],
    )

    document.add_heading("14. Limits of the Study", level=1)
    limits = [
        "The data has only 20 sample records.",
        "The records are made for learning and are not real student records.",
        "The model uses only study hours.",
        "A line may not show all real student cases.",
        "The model is less safe far outside the sample range.",
        "The work does not prove that study hours alone cause marks.",
        "The model must not be used for real school rank, grade, or action.",
    ]
    for limit in limits:
        add_bullet(document, limit)

    document.add_heading("15. Future Work", level=1)
    future_items = [
        "Use real data with clear consent and safe data care.",
        "Use more records from more than one class.",
        "Add class rate, sleep, past marks, and task score.",
        "Check and clean missing or wrong values.",
        "Use more than one train and test split.",
        "Compare the line model with other regression models.",
        "Make a small web page for easy use.",
        "Show a clear note when a new input is outside the learned data range.",
    ]
    for item in future_items:
        add_bullet(document, item)

    document.add_heading("16. Conclusion", level=1)
    add_body_paragraph(
        document,
        "This paper presented a small study-hours and marks prediction project. "
        "The project used Python, Pandas, Scikit-learn, and Matplotlib. It used "
        "Linear Regression because the output is a number and the method is easy "
        "for a new learner."
    )
    add_body_paragraph(
        document,
        "The model gave an R-squared score of 0.997, a Mean Absolute Error of 1.69 "
        "marks, and a Mean Squared Error of 3.47. It gave 51.30 marks for five study "
        "hours. All ten code tests passed."
    )
    add_body_paragraph(
        document,
        "The result is strong only for the small sample pattern. It is not proof "
        "about all students. The project is best used as a clear first step in "
        "machine learning. It helps a learner understand the full path from data "
        "to a model, from a model to a test, and from a test to a user result."
    )


def add_references_and_appendix(document):
    """Add source links and simple support notes."""

    document.add_heading("17. References", level=1)
    references = [
        (
            "Scikit-learn. LinearRegression guide. "
            "https://scikit-learn.org/stable/modules/generated/"
            "sklearn.linear_model.LinearRegression.html"
        ),
        (
            "Scikit-learn. train_test_split guide. "
            "https://scikit-learn.org/stable/modules/generated/"
            "sklearn.model_selection.train_test_split.html"
        ),
        (
            "Pandas. DataFrame guide. "
            "https://pandas.pydata.org/docs/reference/api/pandas.DataFrame.html"
        ),
        (
            "Matplotlib. Scatter plot guide. "
            "https://matplotlib.org/stable/api/_as_gen/matplotlib.pyplot.scatter.html"
        ),
        (
            "Python. Python language guide. https://docs.python.org/3/"
        ),
        (
            "Study Hours vs Marks Prediction project source code and test output, "
            "checked on June 25, 2026."
        ),
    ]
    for reference in references:
        add_numbered_item(document, reference)

    document.add_heading("Appendix A. Main Program Flow", level=1)
    flow_steps = [
        "The program starts.",
        "The sample data is made.",
        "The data is split into train and test parts.",
        "The Linear Regression model is trained.",
        "The test marks are predicted.",
        "The three model test values are shown.",
        "The learned line is shown.",
        "The user enters study hours.",
        "The input is checked.",
        "The final marks are predicted.",
        "The graph is saved and shown.",
        "The program ends.",
    ]
    for flow_step in flow_steps:
        add_numbered_item(document, flow_step)

    document.add_heading("Appendix B. Simple Viva Questions", level=1)
    viva_rows = [
        ("What is the input?", "Study hours."),
        ("What is the output?", "Predicted marks."),
        ("Why use regression?", "The output is a number."),
        ("Why use Linear Regression?", "It is simple and fits a line pattern."),
        ("What is training data?", "Data used to teach the model."),
        ("What is test data?", "Data used to check the trained model."),
        ("What is R-squared?", "A value that shows how well the line fits the data."),
        ("Why check input?", "It stops wrong values and program errors."),
        ("Is the data real?", "No. It is sample data made for learning."),
        ("Can this judge a student?", "No. It is only a beginner project."),
    ]
    add_table(document, ["Question", "Simple Answer"], viva_rows, widths=[2.6, 3.7])


def build_document():
    """Create and save the final research paper."""

    document = Document()
    configure_styles(document)
    configure_page(document)

    add_cover_page(document)
    add_abstract(document)
    add_contents(document)
    add_main_sections(document)
    add_references_and_appendix(document)

    document.core_properties.title = (
        "Study Hours and Marks Prediction Using Linear Regression"
    )
    document.core_properties.subject = "Beginner-friendly machine learning paper"
    document.core_properties.author = "Suvodeep Roy"
    document.core_properties.keywords = (
        "study hours, marks, prediction, linear regression, machine learning"
    )

    document.save(OUTPUT_FILE)
    print(OUTPUT_FILE)


if __name__ == "__main__":
    build_document()
